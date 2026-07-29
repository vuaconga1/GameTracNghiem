#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Split Lớp 9 “Chuyên sâu NP + từ vựng” UNIT docx → skill files like Lớp 8.

Maps theory + practice into:
  Tuvung_Lop9.docx
  NguPhap_Lop9.docx
  Pronunciation_Lop9.docx
  Writing_Lop9.docx
  Reading_Lop9.docx

Practice exercises are classified by title keywords (numbers differ across units).

Usage:
  py -3 web/scripts/data/split_lop9_chuyensau_docx.py --docx "PDF/UNIT 1 -….docx"
  py -3 web/scripts/data/split_lop9_chuyensau_docx.py --docx UNIT1.docx --docx UNIT2.docx …
"""
from __future__ import annotations

import argparse
import re
import sys
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None  # type: ignore

HERE = Path(__file__).resolve().parent
# web/scripts/data → repo root is parents[2]
DEFAULT_OUT = HERE.parents[2] / "PDF"

SKILL_FILES = {
    "tuvung": "Tuvung_Lop9.docx",
    "nguphap": "NguPhap_Lop9.docx",
    "pronunciation": "Pronunciation_Lop9.docx",
    "writing": "Writing_Lop9.docx",
    "reading": "Reading_Lop9.docx",
}

RE_UNIT = re.compile(r"^UNIT\s*(\d+)\s*[.:\-–]?\s*(.*)$", re.I)
RE_FNAME_UNIT = re.compile(r"UNIT\s*(\d+)", re.I)
RE_SECTION = re.compile(
    r"^(I{1,3}|IV|V|VI)\.\s*(VOCABULARY|WORD\s*FORMATION|GRAMMAR|PRONUNCIATION|PRATICE|PRACTICE)\b",
    re.I,
)
RE_EXERCISE = re.compile(r"^Exercise\s*(\d+)\s*[.:]?\s*(.*)$", re.I)

SECTION_SKILL = {
    "vocabulary": "tuvung",
    "word formation": "writing",
    "grammar": "nguphap",
    "pronunciation": "pronunciation",
    "pratice": None,
    "practice": None,
}

# Title → reading (MC / pronunciation drills / reading comprehension) — L8 Reading style
READING_TITLE_RES = [
    re.compile(p, re.I)
    for p in [
        r"different underlined sound",
        r"pronounced differently",
        r"main stress",
        r"words? with\s*/",
        r"sound\s*/",
        r"cluster\s*/",
        r"starting with\s*(with\s*)?/",
        r"correct word with\s*/",
        r"put them in the correct column",
        r"choose the correct option",
        r"choose the correct options",
        r"choose the underlined part",
        r"choose the correct prepositions",
        r"read the (passage|text)",
        r"match the first half",
        r"true or false",
        r"answer the questions",
        r"decide whether the following statements",
    ]
]

# Explicit writing overrides checked first (more specific production tasks)
WRITING_TITLE_RES = [
    re.compile(p, re.I)
    for p in [
        r"look at the photos?",
        r"complete the sentences",
        r"complete the conversation",
        r"complete the (following )?table",
        r"complete the passage with the correct form",
        r"circle the (correct|odd)",
        r"underline the mistake",
        r"rewrite",
        r"reorder",
        r"write (the )?correct sentences",
        r"write complete sentences",
        r"write down the correct",
        r"put the words in brackets",
        r"form of the word",
        r"fill in each blank with the appropriate form",
        r"match the words\s*/\s*phrases with their explanations",
        r"phrasal verbs in the box",
        r"words you've found",
        r"using (double comparatives|past continuous|wish|first conditional|present perfect|may/should)",
        r"suggested words",
    ]
]


def normalize_section_name(raw: str) -> str:
    return re.sub(r"\s+", " ", raw.strip().lower())


def unit_from_filename(path: Path) -> int | None:
    m = RE_FNAME_UNIT.search(path.name)
    return int(m.group(1)) if m else None


def skill_for_exercise_title(title: str) -> str:
    """Classify practice exercise by title keywords (L8 Reading vs Writing split).

    Reading rules run first so MC stems like
    \"Choose the correct option … to complete the sentences\" are not stolen by
    the Writing \"complete the sentences\" pattern.
    """
    t = title.strip()
    for rx in READING_TITLE_RES:
        if rx.search(t):
            return "reading"
    for rx in WRITING_TITLE_RES:
        if rx.search(t):
            return "writing"
    return "writing"


def para_text(element, doc: Document) -> str:
    return Paragraph(element, doc).text.strip()


def table_lead_text(element, doc: Document) -> str:
    """First non-empty line in table (exercise titles sometimes live only in tables)."""
    tbl = Table(element, doc)
    seen: set[int] = set()
    for row in tbl.rows:
        for cell in row.cells:
            # python-docx repeats merged cells; dedupe by cell element id
            cid = id(cell._tc)
            if cid in seen:
                continue
            seen.add(cid)
            raw = (cell.text or "").strip()
            if not raw:
                continue
            return raw.split("\n")[0].strip()
    return ""


def classify_heading(text: str) -> tuple[str, object] | None:
    """Return ('unit', (num, title)) | ('section', name) | ('exercise', (num, title)) | None."""
    if not text:
        return None
    m = RE_UNIT.match(text)
    if m:
        title = (m.group(2) or "").strip(" .-–")
        return ("unit", (int(m.group(1)), title))
    m = RE_SECTION.match(text)
    if m:
        return ("section", normalize_section_name(m.group(2)))
    m = RE_EXERCISE.match(text)
    if m:
        return ("exercise", (int(m.group(1)), (m.group(2) or "").strip()))
    return None


def ensure_before_sectpr(body, element) -> None:
    children = list(body)
    if children and children[-1].tag == qn("w:sectPr"):
        body.insert(len(children) - 1, element)
    else:
        body.append(element)


def _add_image_bytes(dst_doc: Document, blob: bytes) -> str:
    try:
        new_rid, _image = dst_doc.part.get_or_add_image(BytesIO(blob))
        return new_rid
    except UnrecognizedImageError:
        if PILImage is None:
            raise
        im = PILImage.open(BytesIO(blob))
        buf = BytesIO()
        if im.mode in ("P", "RGBA", "LA"):
            im = im.convert("RGBA")
        else:
            im = im.convert("RGB")
        im.save(buf, format="PNG")
        buf.seek(0)
        new_rid, _image = dst_doc.part.get_or_add_image(buf)
        return new_rid


def remap_images(src_doc: Document, dst_doc: Document, element) -> int:
    count = 0
    for blip in element.findall(".//" + qn("a:blip")):
        old_rid = blip.get(qn("r:embed"))
        if not old_rid or old_rid not in src_doc.part.rels:
            continue
        rel = src_doc.part.rels[old_rid]
        if "image" not in rel.reltype:
            continue
        new_rid = _add_image_bytes(dst_doc, rel.target_part.blob)
        blip.set(qn("r:embed"), new_rid)
        count += 1
    return count


def append_element(src_doc: Document, dst_doc: Document, element) -> int:
    new_el = deepcopy(element)
    n_img = remap_images(src_doc, dst_doc, new_el)
    ensure_before_sectpr(dst_doc.element.body, new_el)
    return n_img


def add_heading_para(dst_doc: Document, text: str) -> None:
    dst_doc.add_paragraph(text)


def iter_body_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag in (qn("w:p"), qn("w:tbl")):
            yield child


def split_one_docx(
    src_path: Path,
    skill_docs: dict[str, Document],
    stats: dict[str, dict],
) -> dict:
    src = Document(str(src_path))
    fname_unit = unit_from_filename(src_path)
    unit_num: int | None = fname_unit
    unit_title = ""
    section: str | None = None
    exercise: int | None = None
    in_practice = False
    current_skill: str | None = None
    warned: list[str] = []
    unit_header_written: set[str] = set()
    exercise_routes: list[str] = []

    def ensure_unit_header(skill: str) -> None:
        if skill in unit_header_written or unit_num is None:
            return
        title = f"UNIT {unit_num}"
        if unit_title:
            title = f"UNIT {unit_num}: {unit_title}"
        add_heading_para(skill_docs[skill], title)
        unit_header_written.add(skill)
        stats[skill]["units"].add(unit_num)

    def route_exercise(ex_num: int, ex_title: str, child) -> None:
        nonlocal exercise, in_practice, current_skill
        exercise = ex_num
        in_practice = True
        current_skill = skill_for_exercise_title(ex_title or f"Exercise {ex_num}")
        ensure_unit_header(current_skill)
        n_img = append_element(src, skill_docs[current_skill], child)
        stats[current_skill]["blocks"] += 1
        stats[current_skill]["images"] += n_img
        stats[current_skill]["exercises"].append(ex_num)
        route = f"U{unit_num} Ex{ex_num}→{current_skill}: {(ex_title or '')[:70]}"
        exercise_routes.append(route)

    for child in iter_body_blocks(src):
        text = ""
        if child.tag == qn("w:p"):
            text = para_text(child, src)
        elif child.tag == qn("w:tbl"):
            # Prefer classifying from table lead when it starts with Exercise N
            lead = table_lead_text(child, src)
            heading = classify_heading(lead) if lead else None
            if heading and heading[0] == "exercise":
                ex_num, ex_title = heading[1]
                route_exercise(ex_num, ex_title, child)
                continue
            text = lead  # for orphan warnings only

        heading = classify_heading(text) if text else None
        if heading and child.tag == qn("w:p"):
            kind, payload = heading
            if kind == "unit":
                body_num, unit_title = payload
                if fname_unit is not None and body_num != fname_unit:
                    warned.append(
                        f"Body says UNIT {body_num} but filename is UNIT {fname_unit} "
                        f"— using filename"
                    )
                    unit_num = fname_unit
                else:
                    unit_num = body_num
                section = None
                exercise = None
                in_practice = False
                current_skill = None
                unit_header_written.clear()
                continue
            if kind == "section":
                section = payload
                exercise = None
                in_practice = section in ("practice", "pratice")
                if in_practice:
                    current_skill = None
                else:
                    current_skill = SECTION_SKILL.get(section)
                    if current_skill:
                        ensure_unit_header(current_skill)
                        n_img = append_element(src, skill_docs[current_skill], child)
                        stats[current_skill]["blocks"] += 1
                        stats[current_skill]["images"] += n_img
                        stats[current_skill]["sections"].append(section)
                    else:
                        warned.append(f"Unknown section: {text!r}")
                continue
            if kind == "exercise":
                ex_num, ex_title = payload
                route_exercise(ex_num, ex_title, child)
                continue

        if current_skill is None:
            if text and child.tag == qn("w:p"):
                warned.append(f"Orphan content (no skill yet): {text[:80]!r}")
            continue

        ensure_unit_header(current_skill)
        n_img = append_element(src, skill_docs[current_skill], child)
        stats[current_skill]["blocks"] += 1
        stats[current_skill]["images"] += n_img

    return {
        "unit": unit_num,
        "title": unit_title,
        "warned": warned,
        "routes": exercise_routes,
        "source": str(src_path),
    }


def empty_stats() -> dict[str, dict]:
    return {
        key: {
            "blocks": 0,
            "images": 0,
            "units": set(),
            "sections": [],
            "exercises": [],
        }
        for key in SKILL_FILES
    }


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Split Lớp 9 Chuyên sâu UNIT docx into skill files"
    )
    parser.add_argument(
        "--docx",
        action="append",
        required=True,
        help="Source UNIT docx (repeatable)",
    )
    parser.add_argument(
        "--out-dir",
        default=str(DEFAULT_OUT),
        help=f"Output directory (default: {DEFAULT_OUT})",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Print per-exercise routing",
    )
    args = parser.parse_args()
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    skill_docs = {key: Document() for key in SKILL_FILES}
    for doc in skill_docs.values():
        body = doc.element.body
        for child in list(body):
            if child.tag == qn("w:p") and not Paragraph(child, doc).text:
                body.remove(child)

    stats = empty_stats()

    for raw in args.docx:
        path = Path(raw)
        if not path.is_file():
            print(f"ERROR: not found: {path}", file=sys.stderr)
            return 1
        print(f"Reading: {path.name}")
        summary = split_one_docx(path, skill_docs, stats)
        for w in summary["warned"]:
            print(f"  WARN: {w}")
        if args.verbose:
            for r in summary["routes"]:
                print(f"  {r}")
        print(
            f"  UNIT {summary['unit']}: {summary['title'] or '(no title)'} "
            f"— {len(summary['routes'])} exercises, {len(summary['warned'])} warnings"
        )

    for key, filename in SKILL_FILES.items():
        out_path = out_dir / filename
        skill_docs[key].save(str(out_path))
        s = stats[key]
        print(
            f"Wrote {out_path.name}: blocks={s['blocks']} images={s['images']} "
            f"units={sorted(s['units'])} exercises={s['exercises']}"
        )

    empty = [SKILL_FILES[k] for k, s in stats.items() if s["blocks"] == 0]
    if empty:
        print(f"NOTE: empty skill files: {', '.join(empty)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
