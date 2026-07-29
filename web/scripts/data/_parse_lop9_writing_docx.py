"""
Parse Writing_Lop9.docx → lop9-writing-content.json

Only extracts exercises that can be turned into online games:
  - grammar  (fill-in-the-blank with word box)
  - quiz     (multiple choice / circle correct)
  - choose_and_circle (pick one of two options in sentence)

Skips: rewrite, reorder, match-with-photo, match-definitions, tables-only.
"""

import json
import re
import sys
from pathlib import Path

import docx

DOCX_PATH = Path(__file__).resolve().parents[3] / "PDF" / "Writing_Lop9.docx"
OUT_PATH = Path(__file__).resolve().parent / "lop9-writing-content.json"

doc = docx.Document(str(DOCX_PATH))

paragraphs = [p.text for p in doc.paragraphs]
tables_raw = []
for t in doc.tables:
    rows = []
    for row in t.rows:
        rows.append([c.text.strip() for c in row.cells])
    tables_raw.append(rows)

# ---------------------------------------------------------------------------
# 1. Split paragraphs by UNIT
# ---------------------------------------------------------------------------

units_paras: dict[int, list[str]] = {}
current_unit = 0
for line in paragraphs:
    m = re.match(r"^UNIT\s+(\d+)", line, re.IGNORECASE)
    if m:
        current_unit = int(m.group(1))
        units_paras[current_unit] = []
    if current_unit > 0:
        units_paras[current_unit].append(line)

# ---------------------------------------------------------------------------
# 2. Map tables to exercises
#    Word-box tables are small (≤3 rows, ≥4 cols) with simple word content.
#    We build a table index per unit by scanning paragraph order vs table order.
# ---------------------------------------------------------------------------

def flatten_word_box(table_rows: list[list[str]]) -> list[str]:
    words = []
    for row in table_rows:
        for cell in row:
            w = cell.strip()
            if w and not w.startswith("_") and not re.match(r"^\d+\.\s*_", w):
                words.append(w)
    return words

# Manually assign word-box table indices to (unit, exercise_number)
# Based on the table dump analysis
WORD_BOX_MAP: dict[int, tuple[int, int]] = {
    2: (1, 5),    # T2 → U1 Ex5
    3: (1, 10),   # T3 → U1 Ex10
    4: (1, 12),   # T4 → U1 Ex12
    6: (2, 4),    # T6 → U2 Ex4
    9: (2, 6),    # T9 → U2 Ex6
    10: (2, 8),   # T10 → U2 Ex8
    13: (3, 4),   # T13 → U3 Ex4
    14: (3, 9),   # T14 → U3 Ex9
    15: (3, 11),  # T15 → U3 Ex11
    16: (3, 12),  # T16 → U3 Ex12
    22: (4, 11),  # T22 → U4 Ex11
    27: (5, 10),  # T27 → U5 Ex10
    28: (5, 12),  # T28 → U5 Ex12
    32: (6, 11),  # T32 → U6 Ex11
}

# Ex4 in each unit uses word box from Ex3 photo labels (Table 1/7/25 etc.)
# Manually provide the word lists from photo exercise tables
EX4_WORD_BOXES: dict[int, list[str]] = {
    1: ["artisan", "police officer", "garbage collector", "delivery man", "electrician", "firefighter"],
    2: ["metro", "concrete jungle", "construction site", "sky train", "itchy eyes", "congested road", "entertainment centre", "dusty"],
    4: ["magnificent", "heritage", "thanks to", "well preserved", "occupied", "generation", "takeaway", "structure", "recognise", "contribute"],
    5: ["dance with local people", "tour the campus", "take photos", "ride a jeep", "see a gong show", "climb a mountain", "go snorkeling", "put up tents"],
    6: ["opportunity", "freedom", "memorise", "replace", "pursue", "take notes", "depend on", "various", "family-oriented", "extended"],
}

word_boxes: dict[tuple[int, int], list[str]] = {}
for tidx, (unit, ex) in WORD_BOX_MAP.items():
    word_boxes[(unit, ex)] = flatten_word_box(tables_raw[tidx])

# ---------------------------------------------------------------------------
# 3. Parse exercises from paragraph text
# ---------------------------------------------------------------------------

def parse_fill_blanks(lines: list[str]) -> list[dict]:
    """Parse numbered fill-in-blank lines: '1. The ______ successfully ...'"""
    items = []
    for line in lines:
        m = re.match(r"^(\d+)\.\s+(.+)", line.strip())
        if not m:
            continue
        text = m.group(2).strip()
        if "___" not in text and "______" not in text:
            continue
        parts = re.split(r"_{2,}", text)
        if len(parts) < 2:
            continue
        prefix = parts[0].rstrip()
        suffix = parts[-1].lstrip() if len(parts) > 1 else ""
        items.append({
            "order": int(m.group(1)),
            "prefix": prefix,
            "suffix": suffix,
        })
    return items

def parse_circle_options(lines: list[str]) -> list[dict]:
    """Parse 'N. sentence (optA / optB) rest' → choose_and_circle items."""
    items = []
    for line in lines:
        m = re.match(r"^(\d+)\.\s+(.+)", line.strip())
        if not m:
            continue
        text = m.group(2)
        opt_match = re.search(r"\(([^)]+?)\s*/\s*([^)]+?)\)", text)
        if not opt_match:
            continue
        opt_a = opt_match.group(1).strip()
        opt_b = opt_match.group(2).strip()
        prompt = text
        items.append({
            "order": int(m.group(1)),
            "prompt": prompt,
            "options": [opt_a, opt_b],
            "answer": opt_a,  # first option is correct by convention in this docx
        })
    return items

def parse_circle_multi_options(lines: list[str]) -> list[dict]:
    """Parse 'N. stem optA / optB / optC' with slashes inside sentence.
    
    Handles two formats:
      - Options in parens: "suggests (eat / eating) more fruits"
      - Options inline: "wishes she has /had / is having a better phone"
    """
    items = []
    for line in lines:
        m = re.match(r"^(\d+)\.\s+(.+)", line.strip())
        if not m:
            continue
        text = m.group(2).strip()

        # Format 1: options inside parentheses (optA / optB)
        opt_match = re.search(r"\(([^)]+?)\)", text)
        if opt_match:
            inner = opt_match.group(1)
            if " / " in inner or " I " in inner:
                opts_raw = re.split(r"\s*/\s*|\s+I\s+", inner)
                opts = [o.strip() for o in opts_raw if o.strip()]
                if len(opts) >= 2:
                    items.append({
                        "order": int(m.group(1)),
                        "question": text,
                        "options": opts,
                        "answer": opts[0],
                    })
                    continue

        # Format 2: options inline with " / " or " /" separator (no parens)
        # e.g. "wishes she has /had / is having a better phone"
        # Options are short (1-3 words), separated by /
        if " / " in text or " /" in text:
            # Split around ALL slashes to find the option region
            parts = re.split(r'\s*/\s*', text)
            if len(parts) >= 3:
                # Middle parts are full options; first part ends with option,
                # last part starts with option.
                # First option = last 1-3 words of parts[0]
                pre_words = parts[0].rstrip().split()
                # Last option = first 1-3 words of parts[-1] (before sentence continues)
                post_words = parts[-1].lstrip().split()

                # Determine option length from middle parts (they ARE the options)
                mid_opts = [p.strip() for p in parts[1:-1]]
                avg_words = max(1, sum(len(o.split()) for o in mid_opts) // len(mid_opts)) if mid_opts else 1

                # First option: take last avg_words from pre
                first_opt = ' '.join(pre_words[-avg_words:]) if pre_words else ''
                # Last option: take first avg_words from post
                last_opt_words_list = post_words[:max(avg_words, 2)]
                # Trim trailing punctuation from last option
                last_opt = ' '.join(last_opt_words_list).rstrip('.,?!')

                opts = [first_opt] + mid_opts + [last_opt]
                opts = [o.strip() for o in opts if o.strip()]
                if len(opts) >= 2:
                    items.append({
                        "order": int(m.group(1)),
                        "question": text,
                        "options": opts,
                        "answer": opts[0],
                    })
    return items

def parse_underline_mistake(lines: list[str]) -> list[dict]:
    """Parse 'N. sentence with mistake' → quiz fill_blank."""
    items = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        m = re.match(r"^(\d+)\.\s+(.+)", line)
        if m:
            text = m.group(2).strip()
            if text and "___" not in text:
                items.append({
                    "order": int(m.group(1)),
                    "question": text,
                })
        i += 1
    return items

def parse_word_form(lines: list[str]) -> list[dict]:
    """Parse 'N. sentence (word) ___' or 'N. sentence ___ (word)'."""
    items = []
    for line in lines:
        m = re.match(r"^(\d+)\.\s+(.+)", line.strip())
        if not m:
            continue
        text = m.group(2).strip()
        # Look for (word) pattern indicating base word
        bracket = re.search(r"\((\w+)\)", text)
        if bracket and "___" in text:
            items.append({
                "order": int(m.group(1)),
                "question": text,
                "base_word": bracket.group(1),
            })
    return items

def parse_conversation_fill(lines: list[str]) -> list[dict]:
    """Parse conversation with (N) ___ blanks."""
    items = []
    for line in lines:
        matches = re.finditer(r"\((\d+)\)\s*_{2,}", line)
        for match in matches:
            num = int(match.group(1))
            items.append({
                "order": num,
                "prefix": line[:match.start()].rstrip(),
                "suffix": line[match.end():].lstrip(),
                "full_line": line.strip(),
            })
    return items

def parse_modal_verb_fill(lines: list[str]) -> list[dict]:
    """Parse 'N. If ..., S ___ verb. (negative/positive)'."""
    items = []
    for line in lines:
        m = re.match(r"^(\d+)\.\s+(.+)", line.strip())
        if not m:
            continue
        text = m.group(2).strip()
        if "___" in text:
            polarity_match = re.search(r"\((positive|negative)\)", text, re.IGNORECASE)
            polarity = polarity_match.group(1) if polarity_match else ""
            clean = re.sub(r"\s*\((positive|negative)\)\s*$", "", text, flags=re.IGNORECASE)
            parts = re.split(r"_{2,}", clean)
            items.append({
                "order": int(m.group(1)),
                "prefix": parts[0].rstrip() if parts else "",
                "suffix": parts[-1].lstrip() if len(parts) > 1 else "",
                "hint": polarity,
            })
    return items

# ---------------------------------------------------------------------------
# 4. Process each unit
# ---------------------------------------------------------------------------

def find_exercise_block(lines: list[str], ex_num: int) -> tuple[int, int]:
    """Find start/end indices for Exercise N."""
    pattern = re.compile(
        rf"^Exercise\s+{ex_num}\b[.:]?\s*", re.IGNORECASE
    )
    start = -1
    for i, line in enumerate(lines):
        if pattern.match(line.strip()):
            start = i
            break
    if start < 0:
        return -1, -1
    end = len(lines)
    next_ex = re.compile(r"^Exercise\s+\d+\b", re.IGNORECASE)
    for i in range(start + 1, len(lines)):
        if next_ex.match(lines[i].strip()):
            end = i
            break
    return start, end

def get_exercise_instruction(lines: list[str], start: int) -> str:
    """Extract the instruction text from the exercise header line."""
    line = lines[start].strip()
    m = re.match(r"^Exercise\s+\d+[.:]?\s*(.*)", line, re.IGNORECASE)
    return m.group(1).strip() if m else ""

all_units: dict[str, list] = {}
skipped: list[dict] = []

for unit_num in sorted(units_paras.keys()):
    lines = units_paras[unit_num]
    items: list[dict] = []

    # --- Ex 7: Circle correct verb form (U6) ---
    s, e = find_exercise_block(lines, 7)
    if s >= 0:
        instruction = get_exercise_instruction(lines, s)
        if "circle" in instruction.lower():
            # U6 Ex7: V + to-inf / V-ing answers
            u6_ex7_answers = {
                1: "eating",       # suggest + V-ing
                2: "going",        # fancy + V-ing
                3: "to study",     # decide + to-inf
                4: "to go",        # agree + to-inf
                5: "using",        # avoid + V-ing
                6: "to return",    # promise + to-inf
                7: "waiting",      # mind + V-ing
                8: "playing",      # enjoy + V-ing
                9: "to adopt",     # want + to-inf
                10: "to bake",     # learn + to-inf
                11: "to organise", # plan + to-inf
                12: "cooking",     # finish + V-ing
                13: "trying",      # recommend + V-ing
                14: "to attend",   # hope + to-inf
                15: "to eat",      # tend + to-inf
            }
            circle_items = parse_circle_multi_options(lines[s:e])
            for ci in circle_items:
                correct = u6_ex7_answers.get(ci["order"], ci["answer"]) if unit_num == 6 else ci["answer"]
                items.append({
                    "game": "quiz",
                    "type": "multiple_choice",
                    "typeLabel": "Circle correct form",
                    "skill": "writing",
                    "exercise": f"U{unit_num} Ex7",
                    "question": ci["question"],
                    "answer": correct,
                    "options": ci["options"],
                    "accept": [],
                    "fillMode": False,
                })

    # --- Ex 9: Circle correct form (U4 - wish + past simple) ---
    s, e = find_exercise_block(lines, 9)
    if s >= 0:
        instruction = get_exercise_instruction(lines, s)
        if "circle" in instruction.lower():
            if unit_num == 4:
                # Hardcoded answers: wish + past simple / could
                u4_ex9_answers = {
                    1: "had",
                    2: "could learn",
                    3: "had",
                    4: "visited",
                    5: "knew",
                    6: "could have",
                    7: "had",
                    8: "got",
                    9: "could travel",
                    10: "became",  # first clause; "competed" for second
                }
                circle_items = parse_circle_multi_options(lines[s:e])
                for ci in circle_items:
                    correct = u4_ex9_answers.get(ci["order"], ci["answer"])
                    items.append({
                        "game": "quiz",
                        "type": "multiple_choice",
                        "typeLabel": "Circle correct form",
                        "skill": "writing",
                        "exercise": f"U{unit_num} Ex9",
                        "question": ci["question"],
                        "answer": correct,
                        "options": ci["options"],
                        "accept": [],
                        "fillMode": False,
                    })
            else:
                circle_items = parse_circle_multi_options(lines[s:e])
                for ci in circle_items:
                    items.append({
                        "game": "quiz",
                        "type": "multiple_choice",
                        "typeLabel": "Circle correct form",
                        "skill": "writing",
                        "exercise": f"U{unit_num} Ex9",
                        "question": ci["question"],
                        "answer": ci["answer"],
                        "options": ci["options"],
                        "accept": [],
                        "fillMode": False,
                    })

    # --- Ex 10: Circle correct word (choose_and_circle) ---
    s, e = find_exercise_block(lines, 10)
    if s >= 0:
        instruction = get_exercise_instruction(lines, s)
        if "circle" in instruction.lower():
            circle_opts = parse_circle_options(lines[s:e])
            if circle_opts:
                items.append({
                    "game": "choose_and_circle",
                    "title": f"U{unit_num} Exercise 10",
                    "instruction": instruction,
                    "items": [{
                        "order": c["order"],
                        "image": "",
                        "prompt": c["prompt"],
                        "options": c["options"],
                        "answer": c["answer"],
                    } for c in circle_opts],
                })

    # --- Ex 11: Circle correct word (choose_and_circle) ---
    s, e = find_exercise_block(lines, 11)
    if s >= 0:
        instruction = get_exercise_instruction(lines, s)
        if "circle" in instruction.lower():
            circle_opts = parse_circle_options(lines[s:e])
            if circle_opts:
                items.append({
                    "game": "choose_and_circle",
                    "title": f"U{unit_num} Exercise 11",
                    "instruction": instruction,
                    "items": [{
                        "order": c["order"],
                        "image": "",
                        "prompt": c["prompt"],
                        "options": c["options"],
                        "answer": c["answer"],
                    } for c in circle_opts],
                })

    all_units[str(unit_num)] = items

# ---------------------------------------------------------------------------
# 5. Output
# ---------------------------------------------------------------------------

# Merge with existing data if present
existing: dict[str, list] = {}
if OUT_PATH.exists():
    with open(OUT_PATH, encoding="utf-8") as f:
        existing = json.load(f).get("units", {})

def items_key(item: dict) -> str:
    """Create a dedup key for an item."""
    if item["game"] == "choose_and_circle":
        return f"c&c:{item['title']}"
    elif item["game"] == "quiz":
        return f"quiz:{item['question'][:60]}"
    elif item["game"] == "grammar":
        return f"gr:{item['prefix'][:40]}|{item.get('suffix','')[:20]}"
    return str(item)

merged: dict[str, list] = {}
for u in sorted(set(list(existing.keys()) + list(all_units.keys())), key=int):
    old = existing.get(u, [])
    new = all_units.get(u, [])
    old_keys = {items_key(it) for it in old}
    combined = list(old)
    for it in new:
        if items_key(it) not in old_keys:
            combined.append(it)
    merged[u] = combined

output = {"units": merged, "skipped": skipped}

with open(OUT_PATH, "w", encoding="utf-8") as f:
    json.dump(output, f, ensure_ascii=False, indent=2)

# Print summary
total = 0
for u, items in all_units.items():
    games: dict[str, int] = {}
    for it in items:
        g = it["game"]
        games[g] = games.get(g, 0) + 1
    total += len(items)
    print(f"Unit {u}: {games} ({len(items)} items)")
print(f"\nTotal: {total} items")
print(f"Written to: {OUT_PATH}")
